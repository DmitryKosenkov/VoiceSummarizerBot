"""Converts Markdown (# / ## / - structure) into a Word document."""
import io
import re

from docx import Document
from docx.text.paragraph import Paragraph

# Matches **bold** or *italic* spans (group 1 = marker, group 2 = content);
# the backreference (\1) ensures ** is only closed by ** and not by a
# single *. `code` spans are matched separately as group 3.
_INLINE_PATTERN = re.compile(r"(\*\*|\*)(.+?)\1|`(.+?)`")


def render_markdown_summary_to_docx(markdown_text: str) -> bytes:
    """Build a .docx file from Markdown text and return its bytes."""
    document = Document()

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 2")
            content = line[3:].strip()
        elif line.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            content = line[2:].strip()
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            content = line[2:].strip()
        else:
            paragraph = document.add_paragraph()
            content = line

        _add_inline_runs(paragraph, content)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Add text as one or more runs, turning **bold**/*italic*/`code`
    markers into real character formatting instead of literal characters.
    """
    pos = 0
    for match in _INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])

        marker, styled_text, code_text = match.group(1), match.group(2), match.group(3)
        if code_text is not None:
            paragraph.add_run(code_text).font.name = "Consolas"
        elif marker == "**":
            paragraph.add_run(styled_text).bold = True
        else:
            paragraph.add_run(styled_text).italic = True

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])
